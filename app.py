import os
from collections import defaultdict
from flask import Flask, render_template, request, jsonify
from docx import Document
import pdfplumber
import io
import re


app = Flask(__name__)

# Define the default keywords and their weights for scoring
default_keywords_weights = {
    
}

def parse_docx(file_data):
    doc = Document(io.BytesIO(file_data))
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return text

def parse_pdf(file_data):
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        text = '\n'.join([page.extract_text() for page in pdf.pages])
    return text

def parse_doc(file_data):
    doc = Document(io.BytesIO(file_data))
    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])
    return text

def parse_resume(file_storage):
    _, file_extension = os.path.splitext(file_storage.filename)
    if file_extension == '.docx':
        return parse_docx(file_storage.read())
    elif file_extension == '.doc':
        return parse_doc(file_storage.read())
    elif file_extension == '.pdf':
        return parse_pdf(file_storage.read())
    else:
        raise ValueError("Unsupported file type. Only .docx, .doc, and .pdf files are supported.")

import re

def calculate_score(text, job_requirements, keywords_weights):
    word_count = defaultdict(int)
    words = re.findall(r'\b\w+\b', text.lower())  # Extract words using regular expression
    for word in words:
        word_count[word] += 1

    total_score = 0

    # Calculate score based on job requirements
    for requirement in job_requirements:
        requirement_score = 0
        requirement_words = requirement.lower().split()
        for word in requirement_words:
            if word in word_count:
                requirement_score += word_count[word]
            elif word in keywords_weights:
                requirement_score += keywords_weights[word]
        total_score += requirement_score

    # Calculate score based on remaining keywords
    for keyword, weight in keywords_weights.items():
        if keyword in word_count:
            total_score += word_count[keyword] * weight

    return total_score

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        try:
            job_requirements = request.form.get('job_requirements', '').split(',')
            keyword_weights = {}
            for key, value in request.form.items():
                if key.startswith('keyword_'):
                    keyword = key.replace('keyword_', '')
                    weight_key = f'weight_{keyword}'
                    keyword_weights[value.lower()] = int(request.form[weight_key])
            resumes = request.files.getlist('resumes')
            parsed_resumes = []

            for resume in resumes:
                try:
                    resume_text = parse_resume(resume)
                    score = calculate_score(resume_text, job_requirements, keyword_weights)
                    parsed_resumes.append({'filename': resume.filename, 'score': score})
                except ValueError as e:
                    print(f"Error parsing resume {resume.filename}: {str(e)}")

            sorted_resumes = sorted(parsed_resumes, key=lambda x: x['score'], reverse=True)
            return render_template('results.html', resumes=sorted_resumes, keywords_weights=keyword_weights)

        except Exception as e:
            return render_template('error.html', error=str(e))

    return render_template('index.html', keywords_weights=default_keywords_weights)

@app.route('/add_keyword', methods=['POST'])
def add_keyword():
    new_keyword = request.form.get('new_keyword')
    new_weight = int(request.form.get('new_weight', 1))
    default_keywords_weights[new_keyword.lower()] = new_weight
    return jsonify({'status': 'success'})

if __name__ == '__main__':
    app.run(debug=True)
